import pdfplumber
import io
import base64
from openai import OpenAI
from app.config.settings import settings
import json
from docx import Document  # For reading DOCX files
import magic  # For file type detection
from docx2pdf import convert  # For converting DOCX to PDF
import tempfile
import os
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

openai_client = OpenAI(api_key=settings.OPENAI_API_KEY)

# Define the comprehensive resume parsing prompt
resume_prompt = """
You are an HR assistant designed to extract structured candidate information from resumes and format it in JSON. 
Given a resume text, extract and classify the following fields:
1. Personal Information:
    - FullName
    - PhoneNumber:
        - Number
        - ISDCode (predict if missing)
        - OriginalNumber
        - FormattedNumber
        - Type (mobile, landline)
        - Location (predict from ISD if missing)
    - Email
    - Linkedin
    - Github
    - Other Links (excluding LinkedIn, GitHub, Email)
    - Country (of residence)
    - Nationalities (predict if missing)
    - Date of Birth or Age
    - Gender
    - Marital Status
    - Languages (leave blank if undetectable)
    - Current Job Title (as stated, not inferred)
2. Alternance Check:
    - If any variation or misspelling of 'Alternance' is found in the latest experience → override Job Title with "Alternance"
3. Jobs (suggest relevant roles based on resume)
    - Return a list of job titles matching the profile
4. Degrees:
    For each degree extract:
        - DegreeName
        - NormalizeDegree
        - Specialization
        - Date
        - CountryOrInstitute
5. Certifications:
    For each cert extract:
        - CertificationName
        - IssuingOrganization
        - IssueDate
6. Hard Skills:
    - Max 20 most relevant technical skills
7. Soft Skills:
    - Max 20 most relevant soft skills
8. Professional Experience:
    Include only relevant experiences based on:
        - Current role or career goal
        - Technical skills/industry
        - Not internships or non-professional roles
    For each:
        - Job Title
        - Company
        - Location
        - Start Date
        - End Date (handle "present", "en cours", etc. as "PRESENT")
        - Duration
        - Responsibilities (original text/bullets)
        - Achievements (with metrics)
        - Tools/Technologies used
        - Team Size
        - Relevance Score (High, Medium, Low, Skip)
9. Projects:
    For each:
        - ProjectName
        - Description
        - TechnologiesUsed (array)
        - Role
        - Period
        - URL (if available)
10. Awards and Publications (Prix et Publications):
    For each:
        - Type (Award/Publication)
        - Title
        - Description
        - Date
        - Publisher/Issuer
        - URL (if available)

For any field not found or undetectable, use empty strings or empty arrays.
Return the result in the following JSON structure:
{
    "CandidateInfo": {
        "FullName": "",
        "PhoneNumber": {
            "Number": "",
            "ISDCode": "",
            "OriginalNumber": "",
            "FormattedNumber": "",
            "Type": "",
            "Location": ""
        },
        "Email": "",
        "Linkedin": "",
        "Github": "",
        "OtherLinks": [],
        "Country": "",
        "Nationalities": [],
        "DateOfBirthOrAge": "",
        "Gender": "",
        "MaritalStatus": "",
        "Languages": [],
        "CurrentJobTitle": ""
    },
    "SuggestedJobs": [],
    "Degrees": [
        {
            "DegreeName": "",
            "NormalizeDegree": "",
            "Specialization": "",
            "Date": "",
            "CountryOrInstitute": ""
        }
    ],
    "Certifications": [
        {
            "CertificationName": "",
            "IssuingOrganization": "",
            "IssueDate": ""
        }
    ],
    "HardSkills": [],
    "SoftSkills": [],
    "ProfessionalExperience": [
        {
            "JobTitle": "",
            "Company": "",
            "Location": "",
            "StartDate": "",
            "EndDate": "",
            "Duration": "",
            "Responsibilities": [],
            "Achievements": [],
            "ToolsAndTechnologies": [],
            "TeamSize": "",
            "RelevanceScore": ""
        }
    ],
    "Projects": [
        {
            "ProjectName": "",
            "Description": "",
            "TechnologiesUsed": [],
            "Role": "",
            "Period": "",
            "URL": ""
        }
    ],
    "AwardsAndPublications": [
        {
            "Type": "",
            "Title": "",
            "Description": "",
            "Date": "",
            "PublisherOrIssuer": "",
            "URL": ""
        }
    ]
}
"""

def detect_file_type(binary_data: bytes) -> str:
    """
    Detect file type from binary data using multiple methods.
    
    Args:
        binary_data (bytes): The binary data of the file
        
    Returns:
        str: MIME type of the detected file
    """
    try:
        # Method 1: Use python-magic to detect file type
        mime_type = magic.from_buffer(binary_data, mime=True)
        logger.info(f"Magic detected file type: {mime_type}")
        return mime_type
    except Exception as e:
        logger.warning(f"Magic detection failed: {str(e)}, falling back to signature detection")
        
        # Method 2: Fallback - check file signatures
        if binary_data.startswith(b'%PDF'):
            logger.info("Detected PDF by signature")
            return 'application/pdf'
        elif binary_data.startswith(b'PK\x03\x04'):
            # Check if it's a DOCX file (ZIP-based format)
            if b'word/' in binary_data[:2048] or b'[Content_Types].xml' in binary_data[:2048]:
                logger.info("Detected DOCX by signature")
                return 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            else:
                logger.info("Detected generic ZIP file")
                return 'application/zip'
        elif binary_data.startswith(b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'):
            # Old MS Office format (DOC)
            logger.info("Detected DOC by signature")
            return 'application/msword'
        else:
            logger.warning("Unknown file type")
            return 'unknown'

def extract_text_from_pdf(binary_data: bytes) -> str:
    """
    Extract text from PDF file using pdfplumber.
    
    Args:
        binary_data (bytes): PDF file binary data
        
    Returns:
        str: Extracted text content
        
    Raises:
        ValueError: If PDF extraction fails
    """
    try:
        pdf_text = ""
        with pdfplumber.open(io.BytesIO(binary_data)) as pdf:
            for page_num, page in enumerate(pdf.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        pdf_text += page_text + "\n"
                    logger.debug(f"Extracted text from page {page_num + 1}")
                except Exception as page_error:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {str(page_error)}")
                    continue
        
        pdf_text = pdf_text.strip()
        logger.info(f"Successfully extracted {len(pdf_text)} characters from PDF")
        
        if not pdf_text:
            raise ValueError("No text could be extracted from the PDF")
            
        return pdf_text
        
    except Exception as e:
        logger.error(f"PDF text extraction failed: {str(e)}")
        raise ValueError(f"Error extracting text from PDF: {str(e)}")

def extract_text_from_docx(binary_data: bytes) -> str:
    """
    Extract text from DOCX file using python-docx.
    
    Args:
        binary_data (bytes): DOCX file binary data
        
    Returns:
        str: Extracted text content
        
    Raises:
        ValueError: If DOCX extraction fails
    """
    try:
        # Create a BytesIO object from binary data
        docx_file = io.BytesIO(binary_data)
        
        # Load the document
        doc = Document(docx_file)
        
        # Extract text from all paragraphs
        text_content = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        # Extract text from headers and footers
        for section in doc.sections:
            # Headers
            header = section.header
            for paragraph in header.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Footers
            footer = section.footer
            for paragraph in footer.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
        
        extracted_text = "\n".join(text_content)
        logger.info(f"Successfully extracted {len(extracted_text)} characters from DOCX")
        
        if not extracted_text:
            raise ValueError("No text could be extracted from the DOCX file")
            
        return extracted_text
    
    except Exception as e:
        logger.error(f"DOCX text extraction failed: {str(e)}")
        raise ValueError(f"Error extracting text from DOCX: {str(e)}")

def convert_docx_to_pdf(binary_data: bytes) -> bytes:
    """
    Convert DOCX file to PDF format.
    
    Args:
        binary_data (bytes): DOCX file binary data
        
    Returns:
        bytes: PDF file binary data
        
    Raises:
        ValueError: If conversion fails
    """
    docx_temp_path = None
    pdf_temp_path = None
    
    try:
        # Create temporary files
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as docx_temp:
            docx_temp.write(binary_data)
            docx_temp_path = docx_temp.name
        
        # Create temporary PDF file path
        pdf_temp_path = docx_temp_path.replace('.docx', '.pdf')
        
        logger.info(f"Converting DOCX to PDF: {docx_temp_path} -> {pdf_temp_path}")
        
        # Convert DOCX to PDF using docx2pdf
        convert(docx_temp_path, pdf_temp_path)
        
        # Check if conversion was successful
        if not os.path.exists(pdf_temp_path):
            raise ValueError("PDF conversion failed - output file not created")
        
        # Read the converted PDF
        with open(pdf_temp_path, 'rb') as pdf_file:
            pdf_binary = pdf_file.read()
        
        if not pdf_binary:
            raise ValueError("PDF conversion failed - empty output file")
        
        logger.info(f"Successfully converted DOCX to PDF: {len(pdf_binary)} bytes")
        return pdf_binary
        
    except Exception as e:
        logger.error(f"DOCX to PDF conversion failed: {str(e)}")
        raise ValueError(f"Error converting DOCX to PDF: {str(e)}")
        
    finally:
        # Clean up temporary files
        try:
            if docx_temp_path and os.path.exists(docx_temp_path):
                os.unlink(docx_temp_path)
                logger.debug(f"Cleaned up temporary DOCX file: {docx_temp_path}")
        except Exception as cleanup_error:
            logger.warning(f"Failed to clean up DOCX temp file: {cleanup_error}")
            
        try:
            if pdf_temp_path and os.path.exists(pdf_temp_path):
                os.unlink(pdf_temp_path)
                logger.debug(f"Cleaned up temporary PDF file: {pdf_temp_path}")
        except Exception as cleanup_error:
            logger.warning(f"Failed to clean up PDF temp file: {cleanup_error}")

def process_text_with_openai(extracted_text: str) -> dict:
    """
    Process extracted text using OpenAI API to parse CV information.
    
    Args:
        extracted_text (str): Text extracted from CV
        
    Returns:
        dict: Parsed CV data in structured format
        
    Raises:
        ValueError: If OpenAI processing fails
    """
    try:
        if not settings.OPENAI_API_KEY:
            logger.error("OpenAI API key is not configured")
            raise ValueError("OpenAI API key is not set")
        
        if not extracted_text.strip():
            logger.error("No text provided for OpenAI processing")
            raise ValueError("No text content to process")
        
        logger.info(f"Processing {len(extracted_text)} characters with OpenAI")
        
        # Call OpenAI API with the comprehensive resume parsing prompt
        response = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": resume_prompt},
                {"role": "user", "content": extracted_text}
            ],
            temperature=0.1,
            max_tokens=4000
        )
        
        # Parse the JSON response
        response_content = response.choices[0].message.content
        
        if not response_content:
            raise ValueError("Empty response from OpenAI")
        
        try:
            result = json.loads(response_content)
        except json.JSONDecodeError as json_error:
            logger.error(f"Failed to parse OpenAI response as JSON: {json_error}")
            logger.error(f"Response content: {response_content}")
            raise ValueError(f"Invalid JSON response from OpenAI: {str(json_error)}")
        
        # Validate that we have the required structure
        if not isinstance(result, dict) or "CandidateInfo" not in result:
            logger.error(f"Invalid response structure: {result}")
            raise ValueError("OpenAI response missing required CandidateInfo section")
        
        # Log successful parsing with statistics
        logger.info("Successfully parsed CV using OpenAI")
        logger.info(f"Professional Experience count: {len(result.get('ProfessionalExperience', []))}")
        logger.info(f"Hard Skills count: {len(result.get('HardSkills', []))}")
        logger.info(f"Soft Skills count: {len(result.get('SoftSkills', []))}")
        logger.info(f"Degrees count: {len(result.get('Degrees', []))}")
        logger.info(f"Certifications count: {len(result.get('Certifications', []))}")
        
        if result.get('ProfessionalExperience', []):
            first_exp = result['ProfessionalExperience'][0]
            logger.debug(f"First experience: {first_exp.get('JobTitle', 'N/A')} at {first_exp.get('Company', 'N/A')}")
        else:
            logger.warning("No ProfessionalExperience data found in the parsed result")
        
        return result
        
    except Exception as openai_error:
        logger.error(f"OpenAI API processing failed: {str(openai_error)}")
        raise ValueError(f"OpenAI processing error: {str(openai_error)}")

def create_fallback_response() -> dict:
    """
    Create a fallback response when CV parsing fails.
    
    Returns:
        dict: Basic candidate info structure with "Not Provided" values
    """
    return {
        "CandidateInfo": {
            "FullName": "Not Provided",
            "PhoneNumber": {
                "Number": "",
                "ISDCode": "",
                "OriginalNumber": "",
                "FormattedNumber": "",
                "Type": "",
                "Location": ""
            },
            "Email": "Not Provided",
            "Linkedin": "",
            "Github": "",
            "OtherLinks": [],
            "Country": "",
            "Nationalities": [],
            "DateOfBirthOrAge": "",
            "Gender": "",
            "MaritalStatus": "",
            "Languages": [],
            "CurrentJobTitle": "Not Provided"
        },
        "SuggestedJobs": [],
        "Degrees": [],
        "Certifications": [],
        "HardSkills": [],
        "SoftSkills": [],
        "ProfessionalExperience": [],
        "Projects": [],
        "AwardsAndPublications": []
    }

def parse_cv(binary_data: bytes) -> tuple[dict, bytes]:
    """
    Main function to parse CV from binary data.
    Supports both PDF and DOCX files. DOCX files are converted to PDF.
    
    Args:
        binary_data (bytes): CV file binary data
        
    Returns:
        tuple[dict, bytes]: (parsed_cv_data, pdf_binary_data)
        - parsed_cv_data: Structured CV information
        - pdf_binary_data: PDF version of the CV (original PDF or converted from DOCX)
    """
    logger.info(f"Starting CV parsing for {len(binary_data)} bytes of data")
    
    try:
        # Step 1: Detect file type
        file_type = detect_file_type(binary_data)
        logger.info(f"Detected file type: {file_type}")
        
        extracted_text = ""
        pdf_binary_data = binary_data  # Default to original data
        
        # Step 2: Extract text and handle file conversion based on type
        if file_type == 'application/pdf':
            logger.info("Processing PDF file")
            try:
                extracted_text = extract_text_from_pdf(binary_data)
                pdf_binary_data = binary_data  # Original is already PDF
            except ValueError as pdf_error:
                logger.error(f"PDF processing failed: {str(pdf_error)}")
                return create_fallback_response(), binary_data
        
        elif file_type in [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword'
        ]:
            logger.info("Processing DOCX/Word file")
            try:
                # Extract text from DOCX
                extracted_text = extract_text_from_docx(binary_data)
                
                # Convert DOCX to PDF for storage and viewing
                pdf_binary_data = convert_docx_to_pdf(binary_data)
                logger.info("Successfully converted DOCX to PDF for storage")
                
            except ValueError as docx_error:
                logger.error(f"DOCX processing failed: {str(docx_error)}")
                return create_fallback_response(), binary_data
        
        else:
            logger.error(f"Unsupported file type: {file_type}")
            return create_fallback_response(), binary_data
        
        # Step 3: Validate extracted text
        if not extracted_text.strip():
            logger.error("No text was extracted from the file")
            return create_fallback_response(), pdf_binary_data
        
        logger.info(f"Successfully extracted {len(extracted_text)} characters")
        
        # Step 4: Process with OpenAI
        try:
            parsed_data = process_text_with_openai(extracted_text)
            logger.info("CV parsing completed successfully")
            return parsed_data, pdf_binary_data
            
        except ValueError as openai_error:
            logger.error(f"OpenAI processing failed: {str(openai_error)}")
            return create_fallback_response(), pdf_binary_data
    
    except Exception as e:
        logger.error(f"Unexpected error in parse_cv: {str(e)}")
        import traceback
        logger.error(f"Full traceback: {traceback.format_exc()}")
        return create_fallback_response(), binary_data

# Utility functions for backward compatibility
def get_supported_file_types() -> list:
    """
    Get list of supported file types.
    
    Returns:
        list: List of supported MIME types
    """
    return [
        'application/pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        'application/msword'
    ]

def validate_file_type(binary_data: bytes) -> bool:
    """
    Validate if the file type is supported.
    
    Args:
        binary_data (bytes): File binary data
        
    Returns:
        bool: True if file type is supported, False otherwise
    """
    try:
        file_type = detect_file_type(binary_data)
        return file_type in get_supported_file_types()
    except Exception:
        return False
